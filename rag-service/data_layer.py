"""
Data Layer — document parsing, semantic chunking, embedding, and vector storage.
Chunk quality directly impacts retrieval accuracy.
Each chunk carries source metadata for traceability.
"""

import hashlib
import re
import uuid
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Dict, Optional, Callable, Generator, Tuple

import numpy as np
from loguru import logger

from config import get_config, ChunkingConfig


# --- Data structures
@dataclass
class Document:
    """Original document (title, content, source path and type)."""
    id: str = field(default_factory=lambda: str(uuid.uuid4()))
    title: str = ""
    content: str = ""
    source_path: str = ""           # original file path
    source_type: str = ""           # txt / md / pdf / docx
    metadata: Dict = field(default_factory=dict)
    created_at: str = ""


@dataclass
class Chunk:
    """Atomic retrieval unit. Each chunk is independently retrievable with source metadata
    for traceability. content_hash enables incremental indexing (change detection)."""
    id: str = field(default_factory=lambda: str(uuid.uuid4()))
    document_id: str = ""
    content: str = ""                           # chunk text
    chunk_index: int = 0                        # position in document
    embedding: Optional[List[float]] = None     # vector (to be filled)
    content_hash: str = ""                      # content hash for incremental indexing
    metadata: Dict = field(default_factory=dict)  # {title, source_path, page, section, ...}

    def __post_init__(self):
        if not self.content_hash and self.content:
            self.content_hash = hashlib.md5(self.content.encode("utf-8")).hexdigest()

    def to_dict(self) -> Dict:
        """Convert to ChromaDB storage format."""
        return {
            "id": self.id,
            "document_id": self.document_id,
            "content": self.content,
            "chunk_index": self.chunk_index,
            "content_hash": self.content_hash,
            **self.metadata,
        }


# --- Document parsers (Strategy Pattern)
class BaseDocumentParser(ABC):
    """Abstract base class for document parsers."""

    @abstractmethod
    def parse(self, file_path: Path) -> Document:
        """Parse a file into a Document object."""
        ...

    @staticmethod
    def supports(extension: str) -> bool:
        """Return True if this parser supports the given file extension."""
        ...


class TextParser(BaseDocumentParser):
    """Plain text parser (.txt, .log, .csv, .json, .xml)."""

    @staticmethod
    def supports(extension: str) -> bool:
        return extension.lower() in {".txt", ".log", ".csv", ".json", ".xml"}

    def parse(self, file_path: Path) -> Document:
        logger.info(f"[TextParser] 解析文本文件: {file_path.name}")
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
        return Document(
            title=file_path.stem,
            content=self._clean_text(content),
            source_path=str(file_path),
            source_type=file_path.suffix.lower(),
            metadata={"encoding": "utf-8", "size_bytes": file_path.stat().st_size},
        )

    def _clean_text(self, text: str) -> str:
        """Remove empty lines, normalize line endings."""
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        lines = [line for line in text.split("\n") if line.strip()]
        return "\n".join(lines)


class MarkdownParser(BaseDocumentParser):
    """Markdown parser. Preserves heading hierarchy as metadata; code blocks kept in original but removed for vector retrieval."""

    @staticmethod
    def supports(extension: str) -> bool:
        return extension.lower() in {".md", ".markdown"}

    def parse(self, file_path: Path) -> Document:
        logger.info(f"[MarkdownParser] 解析 Markdown: {file_path.name}")
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()

        # extract heading structure
        sections = self._extract_headings(content)

        return Document(
            title=file_path.stem,
            content=content,
            source_path=str(file_path),
            source_type="md",
            metadata={"sections": sections, "size_bytes": file_path.stat().st_size},
        )

    def _extract_headings(self, content: str) -> List[Dict]:
        """Extract Markdown heading hierarchy."""
        sections = []
        for line in content.split("\n"):
            match = re.match(r"^(#{1,6})\s+(.+)$", line)
            if match:
                level = len(match.group(1))
                sections.append({"level": level, "title": match.group(2).strip()})
        return sections


class PDFParser(BaseDocumentParser):
    """PDF parser using pypdf with page number preservation."""

    @staticmethod
    def supports(extension: str) -> bool:
        return extension.lower() == ".pdf"

    def parse(self, file_path: Path) -> Document:
        logger.info(f"[PDFParser] 解析 PDF: {file_path.name}")
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("请安装 pypdf: pip install pypdf>=4.0.0")

        reader = PdfReader(str(file_path))
        pages_content = []
        page_metadata = []

        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                pages_content.append(f"[第{i+1}页]\n{text}")
                page_metadata.append({"page": i + 1, "char_count": len(text)})

        content = "\n\n".join(pages_content)
        return Document(
            title=file_path.stem,
            content=content,
            source_path=str(file_path),
            source_type="pdf",
            metadata={
                "total_pages": len(reader.pages),
                "pages_with_content": len(page_metadata),
                "page_info": page_metadata,
            },
        )


class WordParser(BaseDocumentParser):
    """Word (.docx) parser. Extracts paragraphs and tables via python-docx."""

    @staticmethod
    def supports(extension: str) -> bool:
        return extension.lower() in {".docx", ".doc"}

    def parse(self, file_path: Path) -> Document:
        logger.info(f"[WordParser] 解析 Word: {file_path.name}")
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx>=1.1.0")

        doc = DocxDocument(str(file_path))
        parts = []

        # extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                style = para.style.name if para.style else "Normal"
                if "Heading" in style or "heading" in style:
                    parts.append(f"## {para.text}")
                else:
                    parts.append(para.text)

        # extract tables
        for table_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                parts.append(f"\n[表格 {table_idx + 1}]\n" + "\n".join(rows))

        content = "\n\n".join(parts)
        return Document(
            title=file_path.stem,
            content=content,
            source_path=str(file_path),
            source_type="docx",
            metadata={"paragraphs": len(doc.paragraphs), "tables": len(doc.tables)},
        )


# --- Parser registry
class DocumentParserRegistry:
    """Parser registry — auto-matches parser by file extension."""

    _parsers: List[BaseDocumentParser] = [
        TextParser(),
        MarkdownParser(),
        PDFParser(),
        WordParser(),
    ]

    @classmethod
    def get_parser(cls, file_path: Path) -> Optional[BaseDocumentParser]:
        extension = file_path.suffix.lower()
        for parser in cls._parsers:
            if parser.supports(extension):
                return parser
        logger.warning(f"[DocumentParserRegistry] 未找到支持 {extension} 的解析器")
        # fallback: treat as plain text
        return TextParser()

    @classmethod
    def register(cls, parser: BaseDocumentParser):
        """Register a custom parser (inserted at front for priority)."""
        cls._parsers.insert(0, parser)
        logger.info(f"[DocumentParserRegistry] 注册解析器: {parser.__class__.__name__}")


# --- Semantic chunker
class SemanticChunker:
    """Semantic chunker — the core RAG component. Splits at paragraph/sentence boundaries
    first, preserves semantic completeness via overlap, and is Chinese-friendly.
    Recommended: chunk_size=512, chunk_overlap=64 tokens."""

    def __init__(self, config: Optional[ChunkingConfig] = None):
        self.config = config or get_config().chunking

    def chunk_document(self, document: Document) -> List[Chunk]:
        """Chunk a document. Markdown: heading-level splitting. Generic: recursive character splitting."""
        logger.info(
            f"[SemanticChunker] 分块文档: {document.title} "
            f"(chunk_size={self.config.chunk_size}, overlap={self.config.chunk_overlap})"
        )

        if document.source_type == "md":
            chunks = self._chunk_markdown(document)
        else:
            chunks = self._chunk_generic(document)

        # 过滤碎片
        chunks = [
            c for c in chunks
            if len(c.content) >= self.config.min_chunk_size
        ]

        # 限制最大分块数
        if len(chunks) > self.config.max_chunks_per_doc:
            logger.warning(
                f"[SemanticChunker] 文档 {document.title} 分块数 {len(chunks)} "
                f"超过上限 {self.config.max_chunks_per_doc}，将被截断"
            )
            chunks = chunks[:self.config.max_chunks_per_doc]

        logger.info(f"[SemanticChunker] 分块完成: {len(chunks)} 个 chunks")
        return chunks

    def _chunk_markdown(self, document: Document) -> List[Chunk]:
        """Chunk Markdown by heading + paragraph boundaries."""
        lines = document.content.split("\n")
        chunks = []
        current_section = ""
        current_content: List[str] = []
        char_count = 0
        # rough estimate: ~2 chars per token for mixed Chinese/English
        max_chars = self.config.chunk_size * 2

        def flush_chunk():
            nonlocal current_content, char_count
            if not current_content:
                return
            text = "\n".join(current_content)
            chunk = Chunk(
                document_id=document.id,
                content=text,
                chunk_index=len(chunks),
                metadata={
                    "title": document.title,
                    "source_path": document.source_path,
                    "source_type": document.source_type,
                    "section": current_section,
                },
            )
            chunks.append(chunk)
            # keep overlap portion as prefix for next chunk
            if self.config.chunk_overlap > 0:
                overlap_chars = self.config.chunk_overlap * 2
                overlap_text = text[-overlap_chars:] if len(text) > overlap_chars else text
                current_content = [overlap_text] if overlap_text else []
                char_count = len(overlap_text)
            else:
                current_content = []
                char_count = 0

        for line in lines:
            # heading line
            if re.match(r"^#{1,6}\s+", line):
                flush_chunk()
                current_section = re.sub(r"^#+\s+", "", line).strip()
                current_content.append(line)
                char_count += len(line) + 1
                continue

            # empty line = potential paragraph boundary
            if not line.strip():
                if current_content:
                    current_content.append("")
                    char_count += 1
                continue

            current_content.append(line)
            char_count += len(line) + 1

            # chunk_size limit reached
            if char_count >= max_chars:
                flush_chunk()

        flush_chunk()  # flush remaining content
        return chunks

    def _chunk_generic(self, document: Document) -> List[Chunk]:
        """Generic recursive chunking. Tries separators in priority order until each chunk fits within chunk_size."""
        return self._recursive_split(
            text=document.content,
            document_id=document.id,
            document_title=document.title,
            source_path=document.source_path,
            source_type=document.source_type,
            separators=self.config.separators,
        )

    def _recursive_split(
        self,
        text: str,
        document_id: str,
        document_title: str,
        source_path: str,
        source_type: str,
        separators: List[str],
        chunk_index_base: int = 0,
    ) -> List[Chunk]:
        """Core recursive split algorithm."""
        max_chars = self.config.chunk_size * 2  # token to char rough conversion

        # base case: text fits within limit
        if len(text) <= max_chars:
            if not text.strip():
                return []
            return [Chunk(
                document_id=document_id,
                content=text.strip(),
                chunk_index=chunk_index_base,
                metadata={
                    "title": document_title,
                    "source_path": source_path,
                    "source_type": source_type,
                },
            )]

        # try splitting with current separator
        separator = separators[0] if separators else "\n"
        next_separators = separators[1:] if len(separators) > 1 else [""]

        if separator:
            splits = text.split(separator)
        else:
            # last resort: force split by character count
            splits = [text[i:i + max_chars] for i in range(0, len(text), max_chars)]

        chunks: List[Chunk] = []
        current_batch = ""
        current_char_count = 0

        for split in splits:
            split_len = len(split)

            # split exceeds limit → recurse with next-level separator
            if split_len > max_chars:
                # flush current batch first
                if current_batch.strip():
                    chunks.extend(self._recursive_split(
                        current_batch, document_id, document_title,
                        source_path, source_type,
                        next_separators,
                        chunk_index_base=chunk_index_base + len(chunks),
                    ))
                    current_batch = ""
                    current_char_count = 0
                # recursively handle oversize split
                sub_chunks = self._recursive_split(
                    split, document_id, document_title,
                    source_path, source_type,
                    next_separators,
                    chunk_index_base=chunk_index_base + len(chunks),
                )
                chunks.extend(sub_chunks)
                continue

            # adding to batch would exceed limit → flush current batch first
            if current_char_count + split_len > max_chars and current_batch.strip():
                chunks.extend(self._recursive_split(
                    current_batch, document_id, document_title,
                    source_path, source_type,
                    next_separators,
                    chunk_index_base=chunk_index_base + len(chunks),
                ))
                current_batch = split
                current_char_count = split_len
                # include separator
                if separator and separator != "":
                    current_batch += separator
                    current_char_count += len(separator)
            else:
                if current_batch:
                    current_batch += (separator if separator else "")
                    current_char_count += len(separator) if separator else 0
                current_batch += split
                current_char_count += split_len

        # flush final batch
        if current_batch.strip():
            chunks.extend(self._recursive_split(
                current_batch, document_id, document_title,
                source_path, source_type,
                next_separators,
                chunk_index_base=chunk_index_base + len(chunks),
            ))

        # add overlap between adjacent chunks
        if self.config.chunk_overlap > 0 and len(chunks) > 1:
            chunks = self._add_overlap(chunks)

        return chunks

    def _add_overlap(self, chunks: List[Chunk]) -> List[Chunk]:
        """Add overlap text between adjacent chunks."""
        overlap_chars = self.config.chunk_overlap * 2
        for i in range(1, len(chunks)):
            prev_content = chunks[i - 1].content
            if len(prev_content) > overlap_chars:
                # take tail of previous chunk as prefix of current chunk
                overlap_text = prev_content[-overlap_chars:]
                # find first complete sentence boundary
                for sep in ["\n\n", "\n", "。", "！", "？"]:
                    if sep in overlap_text:
                        overlap_text = overlap_text[overlap_text.index(sep) + len(sep):]
                        break
                if overlap_text.strip():
                    chunks[i].content = overlap_text + "\n" + chunks[i].content
        return chunks


# --- Embedding generator
class EmbeddingGenerator:
    """Text embedding. Supports DeepSeek API / OpenAI text-embedding-3-small / local (ChromaDB ONNX)."""

    def __init__(self):
        self.config = get_config().llm
        self._embedding_cache: Dict[str, List[float]] = {}
        self._dimension: Optional[int] = None

    @property
    def dimension(self) -> int:
        """Get embedding dimension (lazy, determined from a test embedding)."""
        if self._dimension is None:
            test_embedding = self.embed("test")
            self._dimension = len(test_embedding)
        return self._dimension

    def embed(self, text: str) -> List[float]:
        """Embed a single text. Built-in cache avoids duplicate API calls."""
        cache_key = hashlib.md5(text.encode("utf-8")).hexdigest()
        if cache_key in self._embedding_cache:
            return self._embedding_cache[cache_key]

        logger.debug(f"[EmbeddingGenerator] 向量化文本 ({len(text)} 字符)")

        if self.config.embedding_provider == "openai":
            embedding = self._embed_openai(text)
        elif self.config.embedding_provider == "local":
            embedding = self._embed_local(text)
        else:
            # DeepSeek 作为兜底（通过对话模型间接获取 embedding）
            embedding = self._embed_deepseek(text)

        self._embedding_cache[cache_key] = embedding
        return embedding

    def embed_batch(self, texts: List[str], batch_size: int = 20) -> List[List[float]]:
        """Batch embed texts to reduce API calls and cost."""
        logger.info(f"[EmbeddingGenerator] 批量向量化 {len(texts)} 条文本")

        all_embeddings = []
        uncached = []
        uncached_indices = []

        # check cache first
        for i, text in enumerate(texts):
            cache_key = hashlib.md5(text.encode("utf-8")).hexdigest()
            if cache_key in self._embedding_cache:
                all_embeddings.append(self._embedding_cache[cache_key])
            else:
                uncached.append(text)
                uncached_indices.append(i)
                all_embeddings.append(None)  # placeholder

        if not uncached:
            return all_embeddings

        # batch API calls
        for batch_start in range(0, len(uncached), batch_size):
            batch = uncached[batch_start:batch_start + batch_size]
            batch_embeddings = self._embed_batch_api(batch)

            for j, embedding in enumerate(batch_embeddings):
                idx = uncached_indices[batch_start + j]
                all_embeddings[idx] = embedding

                cache_key = hashlib.md5(uncached[batch_start + j].encode("utf-8")).hexdigest()
                self._embedding_cache[cache_key] = embedding

        return all_embeddings

    def _embed_openai(self, text: str) -> List[float]:
        """OpenAI Embedding API call."""
        from openai import OpenAI
        client = OpenAI(api_key=self.config.openai_api_key)
        response = client.embeddings.create(
            model=self.config.openai_embedding_model,
            input=text,
        )
        return response.data[0].embedding

    def _embed_batch_api(self, texts: List[str]) -> List[List[float]]:
        """Batch embedding API call."""
        if self.config.embedding_provider == "openai":
            from openai import OpenAI
            client = OpenAI(api_key=self.config.openai_api_key)
            response = client.embeddings.create(
                model=self.config.openai_embedding_model,
                input=texts,
            )
            return [d.embedding for d in response.data]
        else:
            # 非批量模式兜底
            return [self.embed(text) for text in texts]

    def _embed_deepseek(self, text: str) -> List[float]:
        """DeepSeek embedding fallback. DeepSeek has no dedicated Embedding API yet,
        so this method falls back to OpenAI. In production, switch to 'openai' or 'local'."""
        logger.warning(
            "[EmbeddingGenerator] DeepSeek 暂未提供专用 Embedding API，"
            "建议将 embedding_provider 设为 'openai' 或 'local'"
        )
        # fall back to OpenAI embedding
        return self._embed_openai(text)

    def _embed_local(self, text: str) -> List[float]:
        """Local embedding via ChromaDB built-in ONNX model (all-MiniLM-L6-v2).
        No API key or additional dependencies needed."""
        # lazy-load ChromaDB default embedding function
        if not hasattr(self, "_chromadb_ef"):
            from chromadb.utils.embedding_functions import DefaultEmbeddingFunction
            logger.info("[EmbeddingGenerator] 加载 ChromaDB 内置 ONNX Embedding 模型...")
            self._chromadb_ef = DefaultEmbeddingFunction()

        embedding = self._chromadb_ef([text])
        return embedding[0]


# --- Vector store manager
class VectorStoreManager:
    """ChromaDB vector database manager. Supports batch/indexing, incremental indexing,
    document deletion, and multi-collection isolation."""

    def __init__(self, collection_name: str = "deepservice_knowledge"):
        self.collection_name = collection_name
        self.config = get_config()

        import chromadb
        from chromadb.config import Settings

        # persistent storage
        db_path = str(self.config.app.vector_db_dir)
        self.client = chromadb.PersistentClient(
            path=db_path,
            settings=Settings(
                anonymized_telemetry=False,
                allow_reset=True,
            ),
        )

        self.collection = self._get_or_create_collection()
        logger.info(
            f"[VectorStoreManager] 初始化完成: collection={collection_name}, "
            f"chunks={self.collection.count()}"
        )

    def _get_or_create_collection(self):
        """Get or create the ChromaDB collection."""
        try:
            return self.client.get_collection(self.collection_name)
        except Exception:
            logger.info(f"[VectorStoreManager] 创建新 Collection: {self.collection_name}")
            return self.client.create_collection(
                name=self.collection_name,
                metadata={
                    "description": "DeepService 企业智能客服知识库",
                    "created_at": "",
                    "hnsw:space": "cosine",
                },
            )

    def index_document(self, document: Document) -> List[Chunk]:
        """Index a single document: chunk → embed → store. Returns all generated Chunks."""
        logger.info(f"[VectorStoreManager] 索引文档: {document.title}")

        # Step 1: semantic chunking
        chunker = SemanticChunker()
        chunks = chunker.chunk_document(document)

        if not chunks:
            logger.warning(f"[VectorStoreManager] 文档 {document.title} 未生成有效 chunk")
            return []

        # Step 2: embed
        embedder = EmbeddingGenerator()
        texts = [c.content for c in chunks]
        embeddings = embedder.embed_batch(texts)

        for chunk, embedding in zip(chunks, embeddings):
            chunk.embedding = embedding

        # Step 3: store in ChromaDB
        self.collection.add(
            ids=[c.id for c in chunks],
            embeddings=[c.embedding for c in chunks],
            documents=[c.content for c in chunks],
            metadatas=[c.to_dict() for c in chunks],
        )

        logger.info(
            f"[VectorStoreManager] 文档 {document.title} 索引完成: "
            f"{len(chunks)} chunks"
        )
        return chunks

    def index_documents(self, documents: List[Document]) -> Dict[str, int]:
        """Index multiple documents. Returns {"indexed": N, "failed": N, "total_chunks": N}."""
        logger.info(f"[VectorStoreManager] 批量索引 {len(documents)} 个文档")
        total_chunks = 0
        indexed = 0
        failed = 0

        for doc in documents:
            try:
                chunks = self.index_document(doc)
                total_chunks += len(chunks)
                indexed += 1
            except Exception as e:
                logger.error(f"[VectorStoreManager] 索引文档 {doc.title} 失败: {e}")
                failed += 1

        return {"indexed": indexed, "failed": failed, "total_chunks": total_chunks}

    def index_file(self, file_path: Path) -> List[Chunk]:
        """Index a single file: auto-detect parser → index.

        Usage: vs = VectorStoreManager(); chunks = vs.index_file(Path("./knowledge_base/faq.md"))
        """
        # parse document
        parser = DocumentParserRegistry.get_parser(file_path)
        if parser is None:
            raise ValueError(f"无法解析文件: {file_path}（不支持的格式）")

        document = parser.parse(file_path)
        return self.index_document(document)

    def index_directory(self, directory: Path, recursive: bool = True) -> Dict[str, int]:
        """Index all supported documents in a directory (recursive by default)."""
        logger.info(f"[VectorStoreManager] 索引目录: {directory}")
        documents = self._load_documents_from_dir(directory, recursive)
        return self.index_documents(documents)

    def delete_document(self, document_id: str) -> int:
        """Delete a document and all its chunks. Returns deleted chunk count."""
        # query all chunks for this document
        results = self.collection.get(
            where={"document_id": document_id},
            include=["metadatas"],
        )

        if results["ids"]:
            self.collection.delete(ids=results["ids"])
            logger.info(
                f"[VectorStoreManager] 删除文档 {document_id}: "
                f"{len(results['ids'])} chunks"
            )
            return len(results["ids"])
        return 0

    def search_by_vector(
        self,
        query_embedding: List[float],
        top_k: int = 5,
        where_filter: Optional[Dict] = None,
    ) -> List[Dict]:
        """Vector semantic search. Supports metadata filtering via where_filter (e.g., {"source_type": "md"})."""
        results = self.collection.query(
            query_embeddings=[query_embedding],
            n_results=top_k,
            where=where_filter,
            include=["documents", "metadatas", "distances"],
        )

        # normalize return format
        formatted = []
        if results["ids"] and results["ids"][0]:
            for i, chunk_id in enumerate(results["ids"][0]):
                distance = results["distances"][0][i]
                # ChromaDB cosine distance -> similarity (1 - distance)
                similarity = 1.0 - distance
                formatted.append({
                    "id": chunk_id,
                    "content": results["documents"][0][i],
                    "metadata": results["metadatas"][0][i] if results["metadatas"] else {},
                    "similarity": round(similarity, 4),
                })

        return formatted

    def get_collection_stats(self) -> Dict:
        """Get collection statistics."""
        count = self.collection.count()
        return {
            "collection_name": self.collection_name,
            "total_chunks": count,
        }

    def reset_collection(self):
        """Reset the collection (dangerous — dev only)."""
        logger.warning(f"[VectorStoreManager] 重置 Collection: {self.collection_name}")
        try:
            self.client.delete_collection(self.collection_name)
            self.collection = self._get_or_create_collection()
        except Exception as e:
            logger.error(f"[VectorStoreManager] 重置失败: {e}")

    def _load_documents_from_dir(
        self,
        directory: Path,
        recursive: bool = True,
    ) -> List[Document]:
        """Load documents from a directory (with extension filter)."""
        documents = []
        pattern = "**/*" if recursive else "*"
        supported_extensions = {".txt", ".md", ".markdown", ".pdf", ".docx"}

        for file_path in directory.glob(pattern):
            if not file_path.is_file():
                continue
            if file_path.suffix.lower() not in supported_extensions:
                continue
            # skip hidden files
            if file_path.name.startswith("."):
                continue

            try:
                parser = DocumentParserRegistry.get_parser(file_path)
                if parser:
                    doc = parser.parse(file_path)
                    documents.append(doc)
            except Exception as e:
                logger.error(f"[VectorStoreManager] 解析 {file_path} 失败: {e}")

        logger.info(f"[VectorStoreManager] 从目录加载了 {len(documents)} 个文档")
        return documents


# --- Self-check
if __name__ == "__main__":
    """Quick self-check. Run: python data_layer.py"""
    logger.info("DeepService Data Layer — self-check")

    # 1. 测试文档解析
    test_dir = get_config().app.knowledge_dir
    if test_dir.exists() and any(test_dir.iterdir()):
        parser_registry = DocumentParserRegistry
        for f in list(test_dir.glob("*"))[:3]:
            parser = parser_registry.get_parser(f)
            if parser:
                doc = parser.parse(f)
                logger.info(f"解析文档: {doc.title} ({len(doc.content)} 字符)")

    # 2. 测试分块
    test_doc = Document(
        title="测试FAQ",
        content=(
            "# 退换货政策\n\n"
            "## 退换货条件\n\n"
            "1. 自签收之日起7天内，商品未经使用且不影响二次销售，可申请退货。\n"
            "2. 自签收之日起15天内，商品出现质量问题，可申请换货。\n\n"
            "## 退换货流程\n\n"
            "1. 登录账号，进入'我的订单'页面\n"
            "2. 选择需退换货的订单，点击'申请售后'\n"
            "3. 填写退换货原因，上传凭证照片\n"
            "4. 等待客服审核（1-3个工作日）\n"
            "5. 审核通过后，按指引寄回商品\n\n"
            "## 运费说明\n\n"
            "- 质量问题退换货：运费由商家承担\n"
            "- 非质量问题退货：运费由买家承担\n"
            "- 换货运费：商家承担寄回运费\n\n"
            "## 退款时效\n\n"
            "- 审核通过后，退款将在3-7个工作日原路返回\n"
            "- 如超过7个工作日未收到退款，请联系客服\n"
        ),
        source_type="md",
    )

    chunker = SemanticChunker()
    chunks = chunker.chunk_document(test_doc)
    for c in chunks:
        logger.info(f"Chunk {c.chunk_index}: {len(c.content)} 字符 | 来源: {c.metadata.get('section', 'N/A')}")

    logger.info("Data layer self-check complete.")
